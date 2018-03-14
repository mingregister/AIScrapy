from bs4 import BeautifulSoup
from selenium import webdriver
import time
import os
import sys
import re
from selenium.webdriver.common.keys import Keys
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from string import Template

phantomjspath = os.path.abspath(
    r'D:\phantomjs\phantomjs\bin\phantomjs.exe')
chromepath = os.path.abspath(
    r'D:\phantomjs\chromedriver_win32\chromedriver.exe')
firefoxpath = os.path.abspath(
    r'D:\phantomjs\geckodriver_win32\geckodriver.exe')

# number = [2, 3, 4, 5, 8, 9, 10]   # 新增后修改*1*
number = [2]
number_map = {                    # 新增后修改*2*
    '2': "A",
    "3": "B",
    "4": "C",
    "5": "D",
    "8": "E",
    "9": "F",
    "10": "G",
}
app_list = ['A','B','C','D','E','F','G']  # 新增后修改*3*
# app_list = ['A']
title_map = {                                                  # 新增后修改*4*
    'A': "a系统",
    "B": "b系统",
    "C": "c系统",
    "D": "d系统",
    "E": "e系统",
    "F": "f系统",
    "G": "g系统",
}
tiers_dict = {                                                 # 新增后修改*5*
    "2": [2, 3, 13, 14],
    # "2":[],
    "3": [5, 6, 7],
    "4": [9, 29, 30],
    "5": [10],
    "8": [15, 16, 17, 18, 19, 20, 21,22],
    "9": [ 23,24, 25, 28],
    "10": [26, 27],
}
A,B,C,D,E,F,G = {},{},{},{},{},{},{}            # 新增后修改*6*
indexs = ["overview", "webTransactions",
          "databaseCalls", "tiers", "remoteServices"] #,"traces"]
# indexs = ["overview"]

def login(username, password, login_url='http://www.baidu.com/pages/v2/login', phantomjspath=os.path.abspath(r'D:\phantomjs\phantomjs\bin\phantomjs.exe')):
    # driver = webdriver.PhantomJS(executable_path=phantomjspath)
    # driver = webdriver.Chrome(executable_path=chromepath)
    driver = webdriver.Firefox()
    driver.get(login_url)
    driver.maximize_window()
    driver.find_element_by_name("input").send_keys(username)
    driver.find_element_by_name("password").send_keys(password)
    driver.find_element_by_name("submit").click()
    return driver


def get_zhouqi(driver, i):
    try:
        if i in range(10):
            xpath = '//*[@id="container"]/div/div[1]/div[1]/div/div/div/div/ul/li[' + \
                str(i) + ']'
        else:
            print(
                'You have choose the Wrong zhouqi_index[1-9].The default options(1Day) has been choose')
            xpath = '//*[@id="container"]/div/div[1]/div[1]/div/div/div/div/ul/li[6]'
    except:
        xpath = '//*[@id="container"]/div/div[1]/div[1]/div/div/div/div/ul/li[6]'
    driver.get('http://www.baidu.com/ai/#/applications')
    driver.find_element_by_xpath(
        '//*[@id="container"]/div/div[1]/div[1]/div/div').click()
    time.sleep(0.5)
    # time.sleep(1)
    driver.find_element_by_xpath(xpath).click()
    return driver


def assemble_url(i, index):
    if not index == "tiers":
        url = "http://www.baidu.com/ai/#/applications/" + \
            str(i) + "/" + index
    else:
        url = "https://www.baidu.com"
        print('get url failed!!!')
    return url


def click_webTransactions_avg(driver):
    '''
    返回 按web事务页面平均响应时间排序结果，从大到小
    :param driver:
    :return:
    '''
    driver.find_element_by_xpath(
        '//*[@id="container"]/div/div[2]/div[3]/div/div/div[1]/table/thead/tr/th[3]').click()
    time.sleep(0.5)
    driver.find_element_by_xpath(
        '//*[@id="container"]/div/div[2]/div[3]/div/div/div[1]/table/thead/tr/th[3]').click()
    time.sleep(0.5)
    return driver


def click_webTransactions_max(driver):
    '''
    返回 按web事务页面 响应时间-最大值 排序结果，从大到小
    :param driver:
    :return:
    '''
    driver.find_element_by_xpath(
        '//*[@id="container"]/div/div[2]/div[3]/div/div/div[1]/table/thead/tr/th[4]').click()
    time.sleep(0.5)
    return driver


def click_remoteServices_Dsf(driver):
    driver.find_element_by_xpath(
        '//*[@id="container"]/div/div[2]/div[3]/nav/a[2]').click()
    time.sleep(0.5)
    return driver


def close(driver):
    driver.close()


def get_html(driver, url):
    driver.get(url)
    time.sleep(3)  # 强行等待，使用implicityly_wait()可能加载不到js
    # time.sleep(10)  # 网络差的时候，就长点吧
    # driver.implicitly_wait(30)
    html = driver.page_source   # <class 'str'>
    return html


def get_html_sorted_webTransactions_avg(driver, url):
    driver.get(url)
    # time.sleep(5)
    time.sleep(2)
    driver = click_webTransactions_avg(driver)
    html = driver.page_source
    return html


def get_html_sorted_webTransactions_max(driver, url):
    driver.get(url)
    # time.sleep(5)
    time.sleep(2)
    driver = click_webTransactions_max(driver)
    html = driver.page_source
    return html


def get_html_sorted_remoteService_Dsf(driver, url):
    driver.get(url)
    # time.sleep(5)
    time.sleep(2)
    driver = click_remoteServices_Dsf(driver)
    html = driver.page_source
    return html


def soup(html, libtype="html5lib"):
    bsObj = BeautifulSoup(html, libtype)
    # print(bsObj.prettify())
    return bsObj


def chose_dict(i):
    try:
        if i == 2:
            dictionally = A
        elif i == 3:
            dictionally = B
        elif i == 4:
            dictionally = C
        elif i == 5:
            dictionally = D
        elif i == 8:
            dictionally = E
        elif i == 9:
            dictionally = F
        elif i == 10:
            dictionally = G
        else:
            print("APP or dictonally is not define")
        return dictionally
    except:
        str_error = "APP or dictonally is not define"
        return str_error


def handle_overview(bsObj):
    ZONGLAN = []
    pattern = re.compile(r'\d+')
    pattern2 = re.compile(r'\d+\.\d+')
    data_webshiwu = bsObj.findAll('p', attrs={'class': 'text-tip'})
    data_yonghutiyan = bsObj.findAll('td')
    data_zonglangtiaoyong = bsObj.findAll('div', attrs={'class': 'pull-right'})
    webshiwujiankang = data_webshiwu[0].get_text()   # Web事务健康
    tierjiankang = data_webshiwu[1].get_text()       # Tier健康情况
    webshiwurukoujiankong = data_webshiwu[2].get_text()       # Web事务入口健康情况
    match1 = pattern.findall(webshiwujiankang)
    match2 = pattern.findall(tierjiankang)
    match3 = pattern.findall(webshiwurukoujiankong)
    match4_manyidu = data_yonghutiyan[2].get_text().strip()          # 满意度
    match5_kerongren = data_yonghutiyan[6].get_text().strip()        # 可容忍度
    match6_huanman = data_yonghutiyan[10].get_text().strip()         # 缓慢
    match7_shibai = data_yonghutiyan[14].get_text().strip()          # 失败
    tiaoyong_per_min = data_zonglangtiaoyong[2].get_text()           # 每分钟调用次数
    response_avg_time = data_zonglangtiaoyong[3].get_text()          # 平均响应时间
    error_per_min = data_zonglangtiaoyong[4].get_text()              # 每分钟错误次数
    match8 = pattern2.findall(tiaoyong_per_min)
    match9 = pattern2.findall(response_avg_time)
    match10 = pattern2.findall(error_per_min)                        # B是空列表
    if len(match10) == 0:match10 = ['0']                             # 当error_per_min中没有数据时，加一个0，
    ZONGLAN.append(match1)
    ZONGLAN.append(match2)
    ZONGLAN.append(match3)
    ZONGLAN.append(match4_manyidu)
    ZONGLAN.append(match5_kerongren)
    ZONGLAN.append(match6_huanman)
    ZONGLAN.append(match7_shibai)
    ZONGLAN.append(match8)
    ZONGLAN.append(match9)
    ZONGLAN.append(match10)
    return ZONGLAN


def handle_table(bsObj):
    global_list = []
    table = bsObj.findAll(
        "table", {'class': 'table table-fixed table-striped table-hover'})[0]
    rows = table.findAll("tr")
    rows = rows[0:5]
    try:
        for row in rows:
            csvRow = []
            for cell in row.findAll(['td', 'th']):
                csvRow.append(cell.get_text())
            global_list.append(csvRow)
    finally:
        try:
            for i in global_list:
                i[2] = i[2].split('\n')[0]
            return global_list
        except:
            return global_list

def handle_table_head(bsObj):
    global_list_head = []
    table = bsObj.findAll(
        "table", {'class': 'table table-fixed table-sortable'})[0]
    rows = table.findAll("tr")
    rows = rows[0:5]
    try:
        for row in rows:
            csvRow = []
            for cell in row.findAll(['td', 'th']):
                cell_text = cell.get_text()
                # if cell_text == '' or cell_text == '操作':continue
                csvRow.append(cell_text)
            global_list_head.append(csvRow)
        if global_list_head[0][0] == '':
            global_list_head[0].pop(0)
        if global_list_head[0][-1] == '操作':
            global_list_head[0].pop()
        return global_list_head
    except:
        print('can not catch data')


def handle_webTransactions(bsObj):
    click_webTransactions_avg(driver)
    return handle_table(bsObj)


def handle_databaseCalls(bsObj):
    return handle_table(bsObj)


def handle_remoteServices(bsObj):
    return handle_table(bsObj)


def handle_tiererrorinfo(bsObj):
    return handle_table(bsObj)

#########

def get_data(index1,index2):
    data = Total_data[index1][index2]
    return data

def initial_file():
    # 打开文档
    headname = u'监控报告-' + time.strftime('%Y-%m-%d',time.localtime(time.time()))
    document = Document()
    document.add_heading(headname, 0)
    return document

def write_text(document, word, fontsize=12, fontname='微软雅黑', bold=False, italic = False,style = None):
    run = document.add_paragraph(style=style)
    run = run.add_run(word)
    run.font.size = Pt(fontsize)
    run.font.name = fontname         #u'微软雅黑'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), fontname)      # u'微软雅黑')
    run.bold = bold
    run.italic = italic
    return document

def add_heading_1(document,index):
    title = title_map[str(index)]
    return write_text(document,title,fontsize=18, style='Heading 1')

def add_heading_2(word,document):
    return write_text(document,word,fontsize=13, style='Heading 2')

def add_table(document,rows,data_list_head):
    if rows == 0:
        rows = 1
    cells = len(data_list_head)
    table = document.add_table(rows=rows,cols=cells,style = 'Table Grid')
    for cell in range(0,cells):
        table.rows[0].cells[cell].text = data_list_head[cell]
    return table

def add_table_overview_1(document):
    table = document.add_table(rows=4,cols=4,style = 'Table Grid')
    table.rows[0].cells[0].text = '检查项目'
    table.rows[0].cells[1].text = '稳定'
    table.rows[0].cells[2].text = '警告'
    table.rows[0].cells[3].text = '严重'
    table.rows[1].cells[0].text = 'Web事务健康情况'
    table.rows[2].cells[0].text = 'Tier健康情况'
    table.rows[3].cells[0].text = 'Web事务入口健康情况'
    return table

def add_table_overview_2(document):
    table = document.add_table(rows=2,cols=5,style = 'Table Grid')
    table.rows[0].cells[0].text = '检查项目'
    table.rows[0].cells[1].text = '满意'
    table.rows[0].cells[2].text = '可容忍度'
    table.rows[0].cells[3].text = '缓慢'
    table.rows[0].cells[4].text = '失败'
    table.rows[1].cells[0].text = '用户体验'
    return table

def add_table_overview_3(document):
    table = document.add_table(rows=2,cols=4,style = 'Table Grid')
    table.rows[0].cells[0].text = '检查项目'
    table.rows[0].cells[1].text = '每分钟调用次数'
    table.rows[0].cells[2].text = '平均响应时间'
    table.rows[0].cells[3].text = '每分钟错误次数'
    table.rows[1].cells[0].text = '总览调用情况'
    return table

def add_table_webTransactions(document,rows,data_list_head):
    return add_table(document,rows,data_list_head)

def add_table_databaseCalls(document,rows,data_list_head):
    return add_table(document,rows,data_list_head)

def add_table_remoteServices(document,rows,data_list_head):

    return add_table(document,rows,data_list_head)

def add_table_jvm(document,rows):
    table = document.add_table(rows=rows, cols=3, style='Table Grid')
    table.rows[0].cells[0].text = 'Tiers'
    table.rows[0].cells[1].text = 'Heap平均使用率'
    table.rows[0].cells[2].text = 'GC情况'
    return table

def add_table_errorInfo(document,rows):
    table = document.add_table(rows=rows, cols=5, style='Table Grid')
    table.rows[0].cells[0].text = '第一次出现时间'
    table.rows[0].cells[1].text = '最后一次出现时间'
    table.rows[0].cells[2].text = '请求地址'
    table.rows[0].cells[3].text = '消息'
    table.rows[0].cells[4].text = '出现次数'
    return table

def fullfill_table(document,data_list,table):
    row = 1
    cell = 0
    row_max = len(data_list)+1
    cell_max = len(data_list[0])
    for row_i in range(row,row_max):
        for cell_i in range(cell,cell_max):
            table.rows[row_i].cells[cell_i].text = data_list[row_i-1][cell_i]
    # document.add_paragraph(u'\n')
    return document

def fullfill_table_overview_1(document,data_list):
    table = add_table_overview_1(document)
    table.rows[1].cells[1].text = data_list[0][0]
    table.rows[1].cells[2].text = data_list[0][1]
    table.rows[1].cells[3].text = data_list[0][2]
    table.rows[2].cells[1].text = data_list[1][0]
    table.rows[2].cells[2].text = data_list[1][1]
    table.rows[2].cells[3].text = data_list[1][2]
    table.rows[3].cells[1].text = data_list[2][0]
    table.rows[3].cells[2].text = data_list[2][1]
    table.rows[3].cells[3].text = data_list[2][2]
    document.add_paragraph(u'\n')
    return document

def fullfill_table_overview_2(document,data_list):
    table = add_table_overview_2(document)
    table.rows[1].cells[1].text = data_list[3]
    table.rows[1].cells[2].text = data_list[4]
    table.rows[1].cells[3].text = data_list[5]
    table.rows[1].cells[4].text = data_list[6]
    document.add_paragraph(u'\n')
    return document

def fullfill_table_overview_3(document,data_list):
    table = add_table_overview_3(document)
    table.rows[1].cells[1].text = data_list[7][0]
    table.rows[1].cells[2].text = data_list[8][0]
    table.rows[1].cells[3].text = data_list[9][0]
    # document.add_paragraph(u'\n')
    return document

def fullfill_table_webTransactions(document,data_list,table):
    return fullfill_table(document,data_list,table)

def fullfill_table_databaseCalls(document,data_list,table):
    return fullfill_table(document,data_list,table)

def fullfill_table_remoteServices(document,data_list,table):
    return fullfill_table(document,data_list,table)

def fullfill_table_jvm(document,data_list,table):
    line_count = 1
    for key in data_list:
        table.rows[line_count].cells[0].text = key
        table.rows[line_count].cells[1].text = data_list[key]
        table.rows[line_count].cells[2].text = '良好'
        line_count = line_count + 1

def fullfill_table_errorInfo(document,data_list,table):
    return fullfill_table(document,data_list,table)


def save(document):
    filename = u'信息系统报告-' + time.strftime('%Y-%m-%d', time.localtime(time.time())) + u'.docx'
    # filename = time.strftime('%Y-%m-%d', time.localtime(time.time())) + u'.docx'
    document.save(filename)

text_template = Template('''如上总览情况，
1、 $app满意度为$grade，系统$state
2、 重点关注事务：
    1）事务$webEntry响应时间相对较长，最长时间为$respond_time;
3、 应用与数据库$database交互时间为$sql_time
4、 与远程服务$remoteSer交互时间为$remoteSer_time,调用次数为$remote_times
详细事务情况、数据库情况、错误情况及其它详见下列内容
''')

text_gaikuang = '''   监控报告系统目前已监控7个系统，25个节点，分别为：a系统（4个节点），网厅客户关系系统（3个节点），c系统（3个节点），d系统（1个节点），e系统（8个节点），f系统（4个节点），事项目录（2个节点）。
    按照监控平台监控统计数据，信息中心各系统今天总体运行良好。
    以下分别对各系统的web事务、远程服务情况、JVM情况、错误信息各项指标进行详细分析：'''


if __name__ == "__main__":
    Total_data = {}
    print('Starting collect data')
    driver = login('zhanghao', 'mima')
    get_zhouqi(driver, 6)
    for i in number:
        dictionally = chose_dict(i)
        for index in indexs:
            if index == "overview":
                url = assemble_url(i, index)
                print(url)
                html = get_html(driver, url)
                bsObj = soup(html)
                data_overview = handle_overview(bsObj)
                dictionally[index] = data_overview
            elif index == "webTransactions":
                url = assemble_url(i, index)
                print(url)
                # the first running, collect avg_time
                html = get_html_sorted_webTransactions_avg(driver, url)
                bsObj = soup(html)
                data_overview = handle_webTransactions(bsObj)
                data_overview_head = handle_table_head(bsObj)
                for i_list in data_overview:
                # 将一维列表的第一个及最后一个数据弹出后，data_overivew也会跟着改变的。
                    try:
                        i_list.pop()
                        i_list.pop(0)
                    except:
                        pass
                dictionally['webTransactions_avg'] = data_overview
                dictionally['webTransactions_avg_head'] = data_overview_head
                # the second running, collect max_time
                html = get_html_sorted_webTransactions_max(driver, url)
                bsObj = soup(html)
                data_overview = handle_webTransactions(bsObj)
                for i_list in data_overview:
                    try:
                        i_list.pop()
                        i_list.pop(0)
                    except:
                        pass
                dictionally['webTransactions_max'] = data_overview
                dictionally['webTransactions_max_head'] = data_overview_head
            elif index == "databaseCalls":
                url = assemble_url(i, index)
                print(url)
                html = get_html(driver, url)
                bsObj = soup(html)
                data_overview = handle_databaseCalls(bsObj)
                data_overview_head = handle_table_head(bsObj)
                for i_list in data_overview:
                    i_list.pop(0)
                dictionally[index] = data_overview
                dictionally[index + '_head'] = data_overview_head
            elif index == "remoteServices":
                url = assemble_url(i, index)
                print(url)
                html = get_html_sorted_remoteService_Dsf(driver, url)
                bsObj = soup(html)
                data_overview = handle_remoteServices(bsObj)
                data_overview_head = handle_table_head(bsObj)
                for i_list in data_overview:
                    i_list.pop(0)
                dictionally[index] = data_overview
                dictionally[index + '_head'] = data_overview_head
            elif index == 'traces':
                url = assemble_url(i, index)
                print(url)
                html = get_html(driver, url)
                bsObj = soup(html)
                data_overview = handle_remoteServices(bsObj)
                data_overview_head = handle_table_head(bsObj)
                for i_list in data_overview:
                    i_list.pop(0)
                dictionally[index] = data_overview
                dictionally[index + '_head'] = data_overview_head
            elif index == "tiers":
                jvm_or_errorinfo = ['jvm', 'errorInfo']
                # tiers_list = []
                tiers_list = {}
                print(tiers_dict[str(i)])
                for choice in jvm_or_errorinfo:
                    node_dict = {}
                    for node in tiers_dict[str(i)]:
                        tiers_url = 'http://www.baidu.com/ai/#/applications/' + \
                            str(i) + '/tiers/' + str(node) + '/' + choice
                        print(tiers_url)
                        driver.get(tiers_url)
                        driver.refresh()          # 强行刷新一次，否则网页链接是改变了，但是内容是没有改变的。
                        # time.sleep(2)
                        time.sleep(1)
                        html = get_html(driver,tiers_url)
                        bsObj = soup(html)
                        dict_key = bsObj.find('h2', attrs={'class': 'title-tab hover'}).get_text()
                        if choice == 'jvm':
                            data = re.findall(
                                r'\d+.\d+\%', bsObj.find('text', attrs={'x': '10'}).find('tspan').get_text())[0]
                            # node_dict[str(i)+str(node)] = data
                            node_dict[dict_key] = data
                        else:
                            data = handle_tiererrorinfo(bsObj)
                            # node_dict[str(i) + str(node)] = data
                            node_dict[dict_key] = data
                        # tiers_list[choice] = node_dict        # 这里可能需要推前一格
                    tiers_list[choice] = node_dict
                dictionally[index] = tiers_list
            else:
                errorstr = "No index match, Please check you URL"
                sys.exit()
        # Total_data.append(dictionally)
        Total_data[number_map[str(i)]] = dictionally
    close(driver)
    print(Total_data)
    print('Collect data finish!!!')
    time.sleep(3600)
